import os
import re
import requests
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_docx_from_md(md_path, docx_path):
    document = Document()
    
    # Set up styles
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    try:
        with open(md_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
    except FileNotFoundError:
        print(f"File not found: {md_path}")
        return

    in_code_block = False

    for line in lines:
        line = line.strip()
        
        # Handle Page Breaks
        if '<div style="page-break-after: always;"></div>' in line:
            document.add_page_break()
            continue
            
        # Handle Images ![Alt](Url)
        img_match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
        if img_match:
            alt_text = img_match.group(1)
            url = img_match.group(2)
            try:
                # Basic check if it's a URL
                if url.startswith('http'):
                    response = requests.get(url, timeout=10)
                    if response.status_code == 200:
                        image_stream = BytesIO(response.content)
                        document.add_picture(image_stream, width=Inches(6.0))
                        last_paragraph = document.paragraphs[-1] 
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        document.add_paragraph(f"[Image: {alt_text} - Could not download]")
                else:
                    document.add_paragraph(f"[Image: {alt_text}]")
            except Exception as e:
                document.add_paragraph(f"[Image: {alt_text} - Error: {str(e)}]")
            continue

        # Handle Headers
        if line.startswith('# '):
            document.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            document.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            document.add_heading(line[4:], level=3)
        
        # Handle Horizontal Rules
        elif line.startswith('---'):
            document.add_paragraph('_' * 50)
            
        # Handle List Items
        elif line.startswith('* ') or line.startswith('- '):
            p = document.add_paragraph(line[2:], style='List Bullet')
            
        elif line[0:2].isdigit() and line[2:4] == '. ':
             p = document.add_paragraph(line, style='List Number')

        # Handle Code Blocks
        elif line.startswith('```'):
            in_code_block = not in_code_block
            continue
            
        else:
            if in_code_block:
                p = document.add_paragraph(line)
                p.style = 'No Spacing'
                font = p.style.font
                font.name = 'Courier New'
                font.size = Pt(10)
            elif line:
                # Basic Bold Parsing **text**
                p = document.add_paragraph()
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        p.add_run(part)

    document.save(docx_path)
    print(f"Successfully created {docx_path}")

if __name__ == "__main__":
    create_docx_from_md('PROJECT_REPORT.md', 'Civic_Report_Project.docx')
